from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "BACKEND_DOCUMENTATION.md"
OUT_DIR = ROOT / "docs"
DOCX_OUT = OUT_DIR / "Signal_Trackers_Backend_KT_Documentation.docx"
HTML_OUT = OUT_DIR / "Signal_Trackers_Backend_KT_Documentation.html"


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_inline_markdown(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def add_docx_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        document.add_heading(strip_inline_markdown(heading.group(2)), level=level)
        return

    if stripped.startswith("- "):
        document.add_paragraph(strip_inline_markdown(stripped[2:]), style="List Bullet")
        return

    numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
    if numbered:
        document.add_paragraph(strip_inline_markdown(numbered.group(1)), style="List Number")
        return

    document.add_paragraph(strip_inline_markdown(stripped))


def build_docx(markdown: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_heading("Signal Trackers Backend KT Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = document.add_paragraph("Generated from BACKEND_DOCUMENTATION.md")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                if code_buffer:
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run("\n".join(code_buffer))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = split_table_row(line)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1

            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header

            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[: len(headers)]):
                    cells[idx].text = value
            continue

        add_docx_paragraph(document, line)
        i += 1

    document.save(DOCX_OUT)


def markdown_to_html(markdown: str) -> str:
    lines = markdown.splitlines()
    body: list[str] = []
    i = 0
    in_ul = False
    in_ol = False
    in_code = False
    code_buffer: list[str] = []

    def close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            body.append("</ul>")
            in_ul = False
        if in_ol:
            body.append("</ol>")
            in_ol = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                close_lists()
                in_code = True
                code_buffer = []
            else:
                in_code = False
                body.append("<pre><code>" + html.escape("\n".join(code_buffer)) + "</code></pre>")
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            close_lists()
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            close_lists()
            headers = split_table_row(stripped)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            body.append("<table><thead><tr>")
            for header in headers:
                body.append(f"<th>{html.escape(header)}</th>")
            body.append("</tr></thead><tbody>")
            for row in rows:
                body.append("<tr>")
                for value in row[: len(headers)]:
                    body.append(f"<td>{html.escape(value)}</td>")
                body.append("</tr>")
            body.append("</tbody></table>")
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            close_lists()
            level = len(heading.group(1))
            body.append(f"<h{level}>{html.escape(strip_inline_markdown(heading.group(2)))}</h{level}>")
            i += 1
            continue

        if stripped.startswith("- "):
            if not in_ul:
                close_lists()
                body.append("<ul>")
                in_ul = True
            body.append(f"<li>{html.escape(strip_inline_markdown(stripped[2:]))}</li>")
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if not in_ol:
                close_lists()
                body.append("<ol>")
                in_ol = True
            body.append(f"<li>{html.escape(strip_inline_markdown(numbered.group(1)))}</li>")
            i += 1
            continue

        close_lists()
        body.append(f"<p>{html.escape(strip_inline_markdown(stripped))}</p>")
        i += 1

    close_lists()

    return """<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Signal Trackers Backend KT Documentation</title>
  <style>
    body { font-family: Arial, sans-serif; line-height: 1.5; margin: 40px; color: #1f2933; }
    h1 { text-align: center; border-bottom: 2px solid #253858; padding-bottom: 12px; }
    h2 { margin-top: 34px; color: #253858; border-bottom: 1px solid #d8dee9; padding-bottom: 6px; }
    h3 { margin-top: 24px; color: #334e68; }
    table { border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 13px; }
    th, td { border: 1px solid #c8d1dc; padding: 8px; vertical-align: top; }
    th { background: #edf2f7; text-align: left; }
    pre { background: #111827; color: #f8fafc; padding: 14px; border-radius: 6px; overflow-x: auto; }
    code { font-family: Consolas, monospace; }
    @media print { body { margin: 18mm; } h2 { page-break-after: avoid; } table { page-break-inside: avoid; } }
  </style>
</head>
<body>
""" + "\n".join(body) + """
</body>
</html>
"""


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    markdown = SOURCE.read_text(encoding="utf-8")
    build_docx(markdown)
    HTML_OUT.write_text(markdown_to_html(markdown), encoding="utf-8")
    print(DOCX_OUT)
    print(HTML_OUT)


if __name__ == "__main__":
    main()
